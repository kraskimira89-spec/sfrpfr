"""Render the landing-page implementation specification as a DOCX file."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "specs" / "10-landing-audit-and-implementation.md"
OUTPUT = ROOT / "docs" / "specs" / "10-landing-audit-and-implementation.docx"
ORDERED_ITEM_RE = re.compile(r"^\d+\.\s+(?P<text>.+)$")


def strip_inline_md(text: str) -> str:
    """Убрать простой inline-markdown (** и `) для plain DOCX-текста."""
    return text.replace("**", "").replace("`", "")


def _dbg(hypothesis_id: str, location: str, message: str, data: dict) -> None:
    # #region agent log
    import json
    import time

    log_path = ROOT / "debug-9e155a.log"
    payload = {
        "sessionId": "9e155a",
        "runId": "pre-fix",
        "hypothesisId": hypothesis_id,
        "location": location,
        "message": message,
        "data": data,
        "timestamp": int(time.time() * 1000),
    }
    with log_path.open("a", encoding="utf-8") as fh:
        fh.write(json.dumps(payload, ensure_ascii=False) + "\n")
    # #endregion


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value


def ordered_item_text(line: str) -> str | None:
    """Return ordered-list content for any Markdown numeric prefix."""
    match = ORDERED_ITEM_RE.match(line)
    return match.group("text") if match else None


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table(document, table_rows)
            table_rows = []

    for line in lines:
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            paragraph = document.add_paragraph()
            paragraph.style = document.styles["Normal"]
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue

        if line.startswith("|") and line.endswith("|"):
            parts = [item.strip() for item in line.strip("|").split("|")]
            if all(set(item) <= {"-", ":", " "} for item in parts):
                continue
            # #region agent log
            if any("**" in p or "`" in p for p in parts):
                _dbg(
                    "H3",
                    "render_landing_tz_docx.py:table",
                    "table row with md",
                    {"parts": [p[:60] for p in parts]},
                )
            # #endregion
            table_rows.append(parts)
            continue

        flush_table()

        if not line:
            continue
        if line.startswith("# "):
            raw = line[2:]
            # #region agent log
            _dbg(
                "H5",
                "render_landing_tz_docx.py:heading",
                "heading branch",
                {"raw": raw[:120], "has_md": ("**" in raw or "`" in raw)},
            )
            # #endregion
            document.add_heading(raw, level=0)
        elif line.startswith("## "):
            raw = line[3:]
            # #region agent log
            _dbg(
                "H5",
                "render_landing_tz_docx.py:heading2",
                "heading2 branch",
                {"raw": raw[:120], "has_md": ("**" in raw or "`" in raw)},
            )
            # #endregion
            document.add_heading(raw, level=1)
        elif line.startswith("### "):
            raw = line[4:]
            # #region agent log
            _dbg(
                "H5",
                "render_landing_tz_docx.py:heading3",
                "heading3 branch",
                {"raw": raw[:120], "has_md": ("**" in raw or "`" in raw)},
            )
            # #endregion
            document.add_heading(raw, level=2)
        elif line.startswith("- "):
            raw = line[2:]
            cleaned = strip_inline_md(raw)
            # #region agent log
            _dbg(
                "H1",
                "render_landing_tz_docx.py:bullet",
                "bullet branch",
                {"raw": raw[:120], "cleaned": cleaned[:120], "stripped": raw != cleaned},
            )
            # #endregion
            document.add_paragraph(cleaned, style="List Bullet")
        elif ordered_text := ordered_item_text(line):
            cleaned = strip_inline_md(ordered_text)
            # #region agent log
            _dbg(
                "H1",
                "render_landing_tz_docx.py:ordered",
                "ordered branch",
                {
                    "raw": ordered_text[:120],
                    "cleaned": cleaned[:120],
                    "stripped": ordered_text != cleaned,
                    "branch_uses_strip": True,
                },
            )
            # #endregion
            document.add_paragraph(cleaned, style="List Number")
        elif line.startswith("**") and line.endswith("**"):
            document.add_paragraph(strip_inline_md(line))
        else:
            cleaned = strip_inline_md(line)
            # #region agent log
            if "**" in line or "`" in line:
                _dbg(
                    "H2",
                    "render_landing_tz_docx.py:else",
                    "else paragraph with md",
                    {"raw": line[:120], "cleaned": cleaned[:120]},
                )
            # #endregion
            document.add_paragraph(cleaned)

    flush_table()
    document.add_page_break()
    document.add_paragraph("Конец документа")
    document.save(OUTPUT)
    # #region agent log
    try:
        from zipfile import ZipFile
        import re as _re

        with ZipFile(OUTPUT) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        stars = len(_re.findall(r"\*\*[^*]+\*\*", xml))
        backticks = xml.count("`")
        sample = ""
        if "Главная" in xml:
            i = xml.index("Главная")
            sample = xml[max(0, i - 20) : i + 40]
        _dbg(
            "H4",
            "render_landing_tz_docx.py:output",
            "docx content scan",
            {
                "literal_bold_md_count": stars,
                "backtick_count": backticks,
                "glavnaya_sample": sample[:80],
                "output": str(OUTPUT),
            },
        )
    except Exception as exc:  # noqa: BLE001
        _dbg("H4", "render_landing_tz_docx.py:output", "docx scan failed", {"error": type(exc).__name__})
    # #endregion
    print(OUTPUT)

if __name__ == "__main__":
    main()
